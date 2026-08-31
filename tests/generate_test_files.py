import os
import shutil
import docx
from PIL import Image, ImageDraw, ImageFont

def generate_sample_data():
    base_dir = r"c:\file finder\sample_test_data"
    source_dir = os.path.join(base_dir, "Source_Images")
    dest_dir = os.path.join(base_dir, "Destination_Folder")
    
    os.makedirs(source_dir, exist_ok=True)
    os.makedirs(dest_dir, exist_ok=True)

    # 1. Generate 5 Sample Images in Source Folder
    images_info = [
        ("product_101.jpg", (220, 100, 100), "Product 101"),
        ("product_102.png", (100, 180, 100), "Product 102"),
        ("IMG_2026.png", (100, 120, 220), "IMG 2026"),
        ("banner_logo.png", (220, 180, 80), "Banner Logo"),
        ("unrelated_photo.jpg", (180, 100, 220), "Unrelated Photo")
    ]

    for filename, color, text in images_info:
        img_path = os.path.join(source_dir, filename)
        img = Image.new('RGB', (300, 200), color=color)
        draw = ImageDraw.Draw(img)
        # Draw decorative border and text
        draw.rectangle([(10, 10), (290, 190)], outline=(255, 255, 255), width=3)
        draw.text((30, 85), text, fill=(255, 255, 255))
        img.save(img_path)

    # 2. Create Sample Word Document (.docx)
    doc_path = os.path.join(base_dir, "Sample_Word_Document.docx")
    doc = docx.Document()
    
    doc.add_heading("Sample Image Extraction List", level=1)
    doc.add_paragraph("This is a sample document containing image names and IDs to test the Word Image Matcher software.")
    
    p1 = doc.add_paragraph()
    p1.add_run("1. product_101.jpg").bold = True
    
    p2 = doc.add_paragraph()
    p2.add_run("2. Image ID: product_102")
    
    doc.add_heading("Table of Images to Fetch:", level=2)
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = 'Table Grid'
    
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = "Item Name"
    hdr_cells[1].text = "Image Reference / ID"
    
    r1_cells = tbl.rows[1].cells
    r1_cells[0].text = "Main Banner Image"
    r1_cells[1].text = "banner_logo"
    
    r2_cells = tbl.rows[2].cells
    r2_cells[0].text = "Photo Gallery 2026"
    r2_cells[1].text = "IMG_2026"

    p3 = doc.add_paragraph()
    p3.add_run("Missing Image Reference Test: missing_item_999.png")

    doc.save(doc_path)
    
    print(f"[+] Sample Word Document created: {doc_path}")
    print(f"[+] Sample Source Folder created: {source_dir} ({len(images_info)} images)")
    print(f"[+] Destination Folder created: {dest_dir} (Empty)")

if __name__ == "__main__":
    generate_sample_data()
