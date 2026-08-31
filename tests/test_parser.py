import unittest
import os
import tempfile
import docx
from src.docx_parser import DocxParser

class TestDocxParser(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()
        self.docx_path = os.path.join(self.temp_dir, "test_doc.docx")

        doc = docx.Document()
        doc.add_paragraph("1. IMG_001.jpg")
        doc.add_paragraph("Image ID: IMG_002")
        
        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "photo_abc"
        table.cell(0, 1).text = "photo_xyz.png"
        table.cell(1, 0).text = "IMG_003"
        table.cell(1, 1).text = "duplicate_item"

        doc.save(self.docx_path)

    def tearDown(self):
        if os.path.exists(self.docx_path):
            os.remove(self.docx_path)
        if os.path.exists(self.temp_dir):
            os.rmdir(self.temp_dir)

    def test_clean_token(self):
        self.assertEqual(DocxParser.clean_token(" 1.  IMG_001.png "), "IMG_001.png")
        self.assertEqual(DocxParser.clean_token("- photo_test "), "photo_test")
        self.assertEqual(DocxParser.clean_token('"quoted_name"'), "quoted_name")

    def test_extract_items(self):
        items = DocxParser.extract_items(self.docx_path)
        self.assertTrue(len(items) >= 5)

        raw_list = [i['raw'] for i in items]
        self.assertIn("IMG_001.jpg", raw_list)
        self.assertIn("IMG_002", raw_list)
        self.assertIn("photo_abc", raw_list)
        self.assertIn("photo_xyz.png", raw_list)

if __name__ == "__main__":
    unittest.main()
