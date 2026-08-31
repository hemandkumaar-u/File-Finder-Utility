import re
import os
from typing import List, Dict, Any
import docx

class DocxParser:
    """Parses Word documents (.docx) to extract image names, IDs, and filenames."""

    @staticmethod
    def clean_token(text: str) -> str:
        """Strip surrounding spaces, quote marks, and bullet markers."""
        if not text:
            return ""
        # Remove common bullet points or numbering prefixes like "1. ", "1)", "a. ", "* ", "- "
        cleaned = re.sub(r'^\s*(?:\d+|[a-zA-Z])[\.\)]\s+', '', text.strip())
        cleaned = re.sub(r'^\s*[\-\*\•\▪\▸]\s*', '', cleaned)
        # Strip quotes and extra spaces
        cleaned = cleaned.strip('\'" \t\r\n')
        return cleaned

    @classmethod
    def extract_items(cls, docx_path: str, extract_tokens_per_line: bool = True) -> List[Dict[str, Any]]:
        """
        Extract candidate image names/IDs from a .docx file.
        Returns a list of dicts: [{'raw': str, 'id': str, 'location': str}]
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Document not found at path: {docx_path}")

        doc = docx.Document(docx_path)
        items: List[Dict[str, Any]] = []
        seen = set()

        COMMON_TABLE_HEADERS = {
            'item name', 'image reference / id', 'image reference', 'image id',
            'filename', 'file name', 'image name', 'description', 'no.', 's.no',
            'sr. no', 'id', 'name', 'sl no', 'sr no', 'reference'
        }

        def add_item(text: str, location: str):
            cleaned = cls.clean_token(text)
            if not cleaned or cleaned.lower() in COMMON_TABLE_HEADERS:
                return

            # Check if text contains labeled prefixes like "Image ID: 123", "File: test.png", "Img: ABC"
            prefix_match = re.search(r'(?:image\s*id|file\s*name|img\s*id|image|file|id)\s*[:\-]\s*(.*)', cleaned, re.IGNORECASE)
            if prefix_match:
                cleaned = prefix_match.group(1).strip()

            # Skip long narrative prose sentences (e.g. > 70 chars with spaces) unless it has a file extension
            has_extension = bool(re.search(r'\.(?:jpg|jpeg|png|gif|bmp|webp|tiff|svg|ico|heic)\b', cleaned, re.IGNORECASE))
            if len(cleaned) > 70 and not has_extension:
                return

            # Split line into tokens if multiple comma/semicolon/newline separated names exist on one line
            lines_or_tokens = [cleaned]
            if extract_tokens_per_line:
                split_tokens = re.split(r'[,;\n\r]+', cleaned)
                if len(split_tokens) > 1:
                    lines_or_tokens = [cls.clean_token(t) for t in split_tokens if cls.clean_token(t)]

            for entry in lines_or_tokens:
                # Ignore entries that look like full sentences (contain spaces and common English stop words without extension)
                if not has_extension and len(entry.split()) > 4 and not re.match(r'^(?:[A-Za-z0-9_-]+)$', entry):
                    continue

                if entry and entry.lower() not in seen:
                    seen.add(entry.lower())
                    items.append({
                        'raw': entry,
                        'id': cls.get_stem_id(entry),
                        'location': location
                    })

        # 1. Extract from Paragraphs
        for idx, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                add_item(text, f"Paragraph {idx + 1}")

        # 2. Extract from Tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        add_item(text, f"Table {t_idx + 1} (Row {r_idx + 1}, Col {c_idx + 1})")

        # 3. Extract from Headers & Footers
        for s_idx, section in enumerate(doc.sections):
            for h_idx, header in enumerate([section.header, section.first_page_header]):
                if header:
                    for p in header.paragraphs:
                        if p.text.strip():
                            add_item(p.text.strip(), f"Header (Section {s_idx + 1})")
            for f_idx, footer in enumerate([section.footer, section.first_page_footer]):
                if footer:
                    for p in footer.paragraphs:
                        if p.text.strip():
                            add_item(p.text.strip(), f"Footer (Section {s_idx + 1})")

        return items

    @staticmethod
    def get_stem_id(text: str) -> str:
        """Returns normalized stem (without file extension or prefix path)."""
        base = os.path.basename(text)
        stem, _ = os.path.splitext(base)
        return stem.strip().lower()
